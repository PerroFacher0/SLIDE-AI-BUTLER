# REDACTOR: AIDEN te escribe DOCUMENTOS completos por voz y te los deja listos en Word.
#
# La herramienta que de verdad le ahorra horas a un estudiante: "escríbeme un ensayo sobre la
# Revolución Industrial y guárdalo" -> AIDEN redacta el documento entero (con su mejor cerebro, el
# experto/Pro), lo guarda como .docx en Documentos\AIDEN y lo abre. Sirve para ensayos, informes,
# cartas, correos, resúmenes, discursos, reseñas... Listo para leer, editar o entregar.

import os
import re
import time

from openai import OpenAI
from secretos import OPENROUTER_API_KEY
from Funciones_Slide.Info.Experto import MODELO_EXPERTO

_cliente = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=OPENROUTER_API_KEY)

_CARPETA = os.path.join(os.path.expanduser("~"), "Documents", "AIDEN")

_ESTILOS = {
    "ensayo": "un ENSAYO académico con introducción, desarrollo argumentado y conclusión",
    "informe": "un INFORME estructurado con secciones (objetivo, desarrollo, resultados, conclusión)",
    "reporte": "un REPORTE claro y estructurado por secciones",
    "carta": "una CARTA formal con encabezado, cuerpo y despedida",
    "correo": "un CORREO electrónico claro y cortés (con asunto sugerido al inicio)",
    "resumen": "un RESUMEN claro y bien organizado de las ideas principales",
    "discurso": "un DISCURSO con gancho inicial, cuerpo y cierre memorable",
    "resena": "una RESEÑA crítica y bien argumentada",
    "articulo": "un ARTÍCULO informativo y ameno con subtítulos",
    "documento": "un DOCUMENTO bien estructurado y claro",
}


def _slug(texto):
    t = re.sub(r"[^\w\s-]", "", str(texto or "").strip().lower())
    t = re.sub(r"\s+", "_", t)
    return (t[:45] or "documento").strip("_")


def _md_a_docx(md, ruta):
    # Convierte el markdown que devuelve el modelo en un .docx con títulos y viñetas. python-docx.
    from docx import Document
    doc = Document()
    for linea in md.splitlines():
        s = linea.rstrip()
        if not s.strip():
            continue
        limpio = s.lstrip("#").strip().replace("**", "").replace("*", "")
        if s.startswith("# "):
            doc.add_heading(limpio, level=0)
        elif s.startswith("## "):
            doc.add_heading(limpio, level=1)
        elif s.startswith("### "):
            doc.add_heading(limpio, level=2)
        elif s.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(s.lstrip()[2:].replace("**", ""), style="List Bullet")
        else:
            doc.add_paragraph(s.replace("**", ""))
    doc.save(ruta)


def _guardar(md, base):
    os.makedirs(_CARPETA, exist_ok=True)
    nombre = f"{base}_{time.strftime('%Y%m%d_%H%M')}"
    # Intenta Word (.docx); si no está python-docx, guarda texto (.md).
    try:
        ruta = os.path.join(_CARPETA, nombre + ".docx")
        _md_a_docx(md, ruta)
        return ruta
    except Exception:
        ruta = os.path.join(_CARPETA, nombre + ".md")
        with open(ruta, "w", encoding="utf-8") as f:
            f.write(md)
        return ruta


def redactar_documento(tema, tipo="documento", abrir=True):
    """HERRAMIENTA: escribe un documento COMPLETO sobre 'tema' y lo guarda en Word (Documentos\\AIDEN),
    abriéndolo. tipo: ensayo | informe | carta | correo | resumen | discurso | reseña | artículo.
    Úsala cuando Marco diga 'escríbeme un ensayo sobre X', 'hazme un informe de Y', 'redacta una carta...'."""
    tema = str(tema or "").strip()
    if not tema:
        return "¿Sobre qué quiere que lo escriba, señor?"
    tkey = str(tipo or "documento").strip().lower()
    tkey = "resena" if tkey in ("reseña", "resena") else tkey
    tkey = "articulo" if tkey in ("artículo", "articulo") else tkey
    estilo = _ESTILOS.get(tkey, _ESTILOS["documento"])

    prompt = (
        f"Escribe {estilo} en ESPAÑOL sobre lo siguiente: {tema}.\n"
        "Requisitos: contenido correcto, bien estructurado y listo para entregar; tono adecuado al "
        "tipo; extensión razonable (varios párrafos, o la que pida el tema). Formatéalo en Markdown: "
        "'# ' para el título, '## ' para las secciones, y viñetas con '- ' donde aporten. NO agregues "
        "comentarios tuyos fuera del documento (nada de 'aquí tienes...'): SOLO el documento."
    )
    try:
        r = _cliente.chat.completions.create(
            model=MODELO_EXPERTO,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2400, temperature=0.6,
        )
        md = (r.choices[0].message.content or "").strip()
    except Exception as e:
        return f"No pude redactarlo, señor: {e}"
    if not md:
        return "No conseguí generar el documento, señor; ¿lo intento de nuevo?"

    ruta = _guardar(md, _slug(tema))
    if abrir:
        try:
            os.startfile(ruta)
        except Exception:
            pass
    try:
        from Nucleo_Slide.Estado_Del_Mundo import registrar_evento
        registrar_evento(f"Redacté un {tkey}: {tema[:50]}", "redactor")
    except Exception:
        pass
    palabras = len(md.split())
    return (f"Listo, señor: escribí su {tkey} sobre '{tema[:50]}' ({palabras} palabras) y se lo abrí. "
            f"Lo guardé en Documentos\\AIDEN. Échele un ojo y me dice si ajusto algo.")
